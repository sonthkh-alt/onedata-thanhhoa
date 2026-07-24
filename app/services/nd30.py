"""Lõi dựng file .docx đúng thể thức văn bản hành chính theo
Nghị định 30/2020/NĐ-CP (CLAUDE.md 8.4).

Đặc tả áp dụng: khổ A4; phông Times New Roman; cỡ chữ nội dung 14;
lề trên/dưới 20 mm, trái 30 mm, phải 15 mm; giãn dòng ~1,4; đầu văn bản
2 cột (cơ quan — quốc hiệu); tên loại "BÁO CÁO" in hoa đậm canh giữa;
kết thúc bằng "./.", khối "Nơi nhận:" và khối chữ ký.
"""

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt
from docx.table import _Cell

PHONG_CHU = "Times New Roman"
CO_CHU_NOI_DUNG = 14
CO_CHU_DAU_TRANG = 13
CO_CHU_NOI_NHAN = 11
GIAN_DONG = 1.4


def _dat_phong(run, co_chu: int, dam: bool = False, nghieng: bool = False) -> None:
    """Đặt phông Times New Roman cho một run (kèm eastAsia để chắc chắn)."""
    run.font.name = PHONG_CHU
    run.font.size = Pt(co_chu)
    run.font.bold = dam
    run.font.italic = nghieng
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), PHONG_CHU)
    rfonts.set(qn("w:cs"), PHONG_CHU)


def _ke_ngang(paragraph, do_rong_cm: float = 0.0) -> None:
    """Vẽ đường kẻ ngắn dưới một đoạn (border-bottom); nếu do_rong_cm > 0
    thì bóp lề hai bên để đường kẻ ngắn lại."""
    if do_rong_cm:
        paragraph.paragraph_format.left_indent = Cm(do_rong_cm)
        paragraph.paragraph_format.right_indent = Cm(do_rong_cm)
    ppr = paragraph._element.get_or_add_pPr()
    pbdr = ppr.makeelement(qn("w:pBdr"), {})
    bottom = ppr.makeelement(
        qn("w:bottom"),
        {
            qn("w:val"): "single",
            qn("w:sz"): "6",
            qn("w:space"): "1",
            qn("w:color"): "000000",
        },
    )
    pbdr.append(bottom)
    ppr.append(pbdr)


def _doan_trong_o(
    o: _Cell,
    van_ban: str,
    co_chu: int = CO_CHU_DAU_TRANG,
    dam: bool = False,
    nghieng: bool = False,
    can_giua: bool = True,
    dung_doan_dau: bool = False,
):
    """Thêm một đoạn văn bản vào ô bảng."""
    doan = (
        o.paragraphs[0]
        if dung_doan_dau and not o.paragraphs[0].runs
        else o.add_paragraph()
    )
    doan.alignment = WD_ALIGN_PARAGRAPH.CENTER if can_giua else WD_ALIGN_PARAGRAPH.LEFT
    doan.paragraph_format.space_after = Pt(0)
    run = doan.add_run(van_ban)
    _dat_phong(run, co_chu, dam=dam, nghieng=nghieng)
    return doan


def tao_van_ban() -> Document:
    """Tạo Document khổ A4, lề chuẩn NĐ30, phông mặc định Times New Roman."""
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    sec.top_margin = Mm(20)
    sec.bottom_margin = Mm(20)
    sec.left_margin = Mm(30)
    sec.right_margin = Mm(15)

    kieu = doc.styles["Normal"]
    kieu.font.name = PHONG_CHU
    kieu.font.size = Pt(CO_CHU_NOI_DUNG)
    kieu.element.rPr.rFonts.set(qn("w:eastAsia"), PHONG_CHU)
    kieu.paragraph_format.line_spacing = GIAN_DONG
    kieu.paragraph_format.space_after = Pt(6)
    return doc


def them_phan_dau(
    doc: Document,
    co_quan_chu_quan: str,
    co_quan_ban_hanh: str,
    so_ky_hieu: str,
    dia_danh: str,
    ngay: str,
) -> None:
    """Đầu văn bản 2 cột: trái = cơ quan; phải = quốc hiệu, tiêu ngữ, ngày."""
    bang = doc.add_table(rows=1, cols=2)
    bang.alignment = WD_TABLE_ALIGNMENT.CENTER
    bang.autofit = False
    bang.columns[0].width = Cm(7.0)
    bang.columns[1].width = Cm(9.5)

    o_trai = bang.cell(0, 0)
    _doan_trong_o(o_trai, co_quan_chu_quan.upper(), dung_doan_dau=True)
    doan_cq = _doan_trong_o(o_trai, co_quan_ban_hanh.upper(), dam=True)
    _ke_ngang(doan_cq, do_rong_cm=1.5)
    _doan_trong_o(o_trai, so_ky_hieu, co_chu=CO_CHU_NOI_DUNG - 1)

    o_phai = bang.cell(0, 1)
    _doan_trong_o(
        o_phai,
        "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM",
        co_chu=CO_CHU_DAU_TRANG - 1,
        dam=True,
        dung_doan_dau=True,
    )
    doan_tieu_ngu = _doan_trong_o(
        o_phai, "Độc lập - Tự do - Hạnh phúc", co_chu=CO_CHU_DAU_TRANG, dam=True
    )
    _ke_ngang(doan_tieu_ngu, do_rong_cm=1.8)
    _doan_trong_o(
        o_phai,
        f"{dia_danh}, {ngay}",
        co_chu=CO_CHU_DAU_TRANG,
        nghieng=True,
    )
    doc.add_paragraph()


def them_ten_loai(doc: Document, ten_loai: str, trich_yeu: str) -> None:
    """Tên loại văn bản (BÁO CÁO) in hoa đậm canh giữa; trích yếu đậm,
    có đường kẻ ngắn phía dưới."""
    doan_ten = doc.add_paragraph()
    doan_ten.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doan_ten.paragraph_format.space_after = Pt(0)
    _dat_phong(doan_ten.add_run(ten_loai.upper()), CO_CHU_NOI_DUNG, dam=True)

    doan_trich_yeu = doc.add_paragraph()
    doan_trich_yeu.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _dat_phong(doan_trich_yeu.add_run(trich_yeu), CO_CHU_NOI_DUNG, dam=True)
    _ke_ngang(doan_trich_yeu, do_rong_cm=5.5)
    doc.add_paragraph()


def them_doan(doc: Document, van_ban: str, dam: bool = False, nghieng: bool = False):
    """Đoạn nội dung: canh đều, thụt đầu dòng 1 cm, giãn dòng 1,4."""
    doan = doc.add_paragraph()
    doan.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doan.paragraph_format.first_line_indent = Cm(1.0)
    doan.paragraph_format.line_spacing = GIAN_DONG
    _dat_phong(doan.add_run(van_ban), CO_CHU_NOI_DUNG, dam=dam, nghieng=nghieng)
    return doan


def them_muc(doc: Document, so_muc: str, tieu_de: str) -> None:
    """Đề mục trong nội dung (ví dụ "1. Tình hình chung") — đậm."""
    them_doan(doc, f"{so_muc} {tieu_de}", dam=True)


def them_ket_thuc(
    doc: Document,
    noi_nhan: list[str],
    chuc_vu_dong_1: str,
    chuc_vu_dong_2: str,
    ho_ten: str,
) -> None:
    """Khối kết thúc: "Nơi nhận:" bên trái; khối chữ ký bên phải."""
    bang = doc.add_table(rows=1, cols=2)
    bang.autofit = False
    bang.columns[0].width = Cm(8.0)
    bang.columns[1].width = Cm(8.5)

    o_trai = bang.cell(0, 0)
    doan_nn = o_trai.paragraphs[0]
    doan_nn.paragraph_format.space_after = Pt(0)
    run_nn = doan_nn.add_run("Nơi nhận:")
    _dat_phong(run_nn, CO_CHU_NOI_NHAN + 1, dam=True, nghieng=True)
    for dong in noi_nhan:
        doan = o_trai.add_paragraph()
        doan.paragraph_format.space_after = Pt(0)
        _dat_phong(doan.add_run(f"- {dong};"), CO_CHU_NOI_NHAN)
    doan_luu = o_trai.add_paragraph()
    doan_luu.paragraph_format.space_after = Pt(0)
    _dat_phong(doan_luu.add_run("- Lưu: VT."), CO_CHU_NOI_NHAN)

    o_phai = bang.cell(0, 1)
    _doan_trong_o(
        o_phai,
        chuc_vu_dong_1.upper(),
        co_chu=CO_CHU_DAU_TRANG,
        dam=True,
        dung_doan_dau=True,
    )
    _doan_trong_o(o_phai, chuc_vu_dong_2.upper(), co_chu=CO_CHU_DAU_TRANG, dam=True)
    for _ in range(3):  # khoảng trống chữ ký
        _doan_trong_o(o_phai, "")
    _doan_trong_o(o_phai, ho_ten, co_chu=CO_CHU_NOI_DUNG, dam=True)


def them_dong_mo_phong(doc: Document) -> None:
    """Dòng bắt buộc ở mọi file demo: tuyên bố dữ liệu mô phỏng."""
    doan = doc.add_paragraph()
    doan.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _dat_phong(
        doan.add_run(
            "Dữ liệu mô phỏng phục vụ trình diễn — "
            "không phải số liệu thống kê chính thức."
        ),
        CO_CHU_NOI_NHAN,
        nghieng=True,
    )
