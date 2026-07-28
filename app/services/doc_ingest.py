"""Tiếp nhận văn bản vào Lớp 1: đọc .docx → toàn văn + siêu dữ liệu →
chia đoạn → chỉ mục FTS5 (CLAUDE.md 9.2).

Trong thực tế đây là luồng tự động từ TD Office; demo mô phỏng bằng tải lên.
Bóc siêu dữ liệu bằng biểu thức chính quy trên phần đầu văn bản; người dùng
được sửa lại trước khi lưu.
"""

import re
from datetime import date, datetime
from pathlib import Path

from docx import Document
from sqlalchemy import text
from sqlalchemy.orm import Session

from app.models import DonVi, VanBan, VanBanDoan

DO_DAI_DOAN_TOI_THIEU = 40  # ký tự — đoạn quá ngắn được gộp vào đoạn sau


def doc_toan_van_docx(duong_dan: str | Path) -> str:
    """Đọc toàn văn một file .docx (đoạn văn + nội dung bảng)."""
    doc = Document(str(duong_dan))
    cac_dong: list[str] = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for bang in doc.tables:
        for hang in bang.rows:
            for o in hang.cells:
                for p in o.paragraphs:
                    if p.text.strip() and p.text.strip() not in cac_dong:
                        cac_dong.append(p.text.strip())
    return "\n".join(cac_dong)


def boc_sieu_du_lieu(toan_van: str, db: Session) -> dict:
    """Cố gắng bóc số, ký hiệu, trích yếu, ngày, cơ quan từ phần đầu văn bản.

    Người dùng luôn được sửa lại trước khi lưu — đây chỉ là điền sẵn.
    """
    dau_van_ban = toan_van[:2000]
    ket_qua: dict = {
        "so": "",
        "ky_hieu": "",
        "loai": "bao_cao",
        "trich_yeu": "",
        "co_quan_id": None,
        "ngay_ban_hanh": None,
    }

    # Số + ký hiệu: "Số: 45/BC-UBND" hoặc "Số:        /BC-UBND"
    khop = re.search(r"Số:\s*(\d+)?\s*/\s*([A-ZĐ][\w-]+)", dau_van_ban)
    if khop:
        ket_qua["so"] = khop.group(1) or ""
        ket_qua["ky_hieu"] = khop.group(2) or ""

    # Loại văn bản theo ký hiệu / tiêu đề
    chu_hoa = dau_van_ban.upper()
    if "BÁO CÁO" in chu_hoa or "BC-" in chu_hoa:
        ket_qua["loai"] = "bao_cao"
    elif "KẾ HOẠCH" in chu_hoa or "KH-" in chu_hoa:
        ket_qua["loai"] = "ke_hoach"
    elif "THÔNG BÁO" in chu_hoa or "TB-" in chu_hoa:
        ket_qua["loai"] = "thong_bao"
    else:
        ket_qua["loai"] = "cong_van"

    # Ngày ban hành: "ngày 24 tháng 7 năm 2026"
    khop = re.search(r"ngày\s+(\d{1,2})\s+tháng\s+(\d{1,2})\s+năm\s+(\d{4})", toan_van)
    if khop:
        try:
            ket_qua["ngay_ban_hanh"] = date(
                int(khop.group(3)), int(khop.group(2)), int(khop.group(1))
            )
        except ValueError:
            pass

    # Trích yếu: dòng sau tiêu đề "BÁO CÁO"/"KẾ HOẠCH"... hoặc dòng "V/v ..."
    cac_dong = [d.strip() for d in toan_van.split("\n") if d.strip()]
    for i, dong in enumerate(cac_dong[:15]):
        if dong.upper() in ("BÁO CÁO", "KẾ HOẠCH", "THÔNG BÁO"):
            if i + 1 < len(cac_dong):
                ket_qua["trich_yeu"] = cac_dong[i + 1][:300]
            break
        if dong.startswith("V/v"):
            ket_qua["trich_yeu"] = dong[3:].strip(" :.,")[:300]
            break

    # Cơ quan ban hành: dò tên đơn vị trong phần đầu văn bản
    for dv in db.query(DonVi).all():
        if dv.ten.upper() in chu_hoa or f"UBND {dv.ten}".upper() in chu_hoa:
            ket_qua["co_quan_id"] = dv.id
            if dv.loai in ("xa", "phuong"):
                break  # ưu tiên khớp được cấp xã
    return ket_qua


def chia_doan(toan_van: str) -> list[str]:
    """Chia toàn văn thành các đoạn đủ dài để tìm kiếm."""
    ket_qua: list[str] = []
    dem: list[str] = []
    for dong in toan_van.split("\n"):
        dong = dong.strip()
        if not dong:
            continue
        dem.append(dong)
        if sum(len(d) for d in dem) >= DO_DAI_DOAN_TOI_THIEU:
            ket_qua.append(" ".join(dem))
            dem = []
    if dem:
        ket_qua.append(" ".join(dem))
    return ket_qua


def luu_van_ban(
    db: Session,
    toan_van: str,
    so: str,
    ky_hieu: str,
    loai: str,
    trich_yeu: str,
    co_quan_id: int | None,
    ngay_ban_hanh: date | None,
    duong_dan_file: str | None,
    mat: bool = False,
) -> VanBan:
    """Lưu văn bản vào Lớp 1: bản ghi + đoạn + chỉ mục FTS5.

    Văn bản `mat=True` KHÔNG được chia đoạn/lập chỉ mục — bị chặn khỏi
    tìm kiếm, AI và trích xuất ngay từ tầng lưu trữ.
    """
    vb = VanBan(
        so=so.strip()[:20],
        ky_hieu=ky_hieu.strip()[:60],
        loai=loai,
        trich_yeu=trich_yeu.strip()[:500] or "(chưa có trích yếu)",
        co_quan_id=co_quan_id,
        ngay_ban_hanh=ngay_ban_hanh,
        duong_dan_file=duong_dan_file,
        toan_van=toan_van,
        mat=mat,
        thoi_diem_tiep_nhan=datetime.now(),
    )
    db.add(vb)
    db.flush()
    if not mat:
        for i, doan in enumerate(chia_doan(toan_van)):
            d = VanBanDoan(van_ban_id=vb.id, thu_tu=i, noi_dung=doan)
            db.add(d)
            db.flush()
            db.execute(
                text(
                    "INSERT INTO van_ban_fts (noi_dung, van_ban_id, doan_id) "
                    "VALUES (:nd, :vb, :d)"
                ),
                {"nd": doan, "vb": vb.id, "d": d.id},
            )
    db.commit()
    return vb
