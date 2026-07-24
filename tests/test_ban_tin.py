"""Kiểm thử máy tham mưu: dự báo, bất thường, bản tin, công văn, hộ chiếu
số liệu, đồng hồ tiết kiệm."""

import io

import pytest
from docx import Document

from app.models import DonVi, GiaTriChiTieu
from app.services import ban_tin
from app.services.kiem_ke import dong_ho_tiet_kiem

MAT_KHAU = "Demo@2026"


def dang_nhap(client, ten_dang_nhap):
    resp = client.post(
        "/dang-nhap", data={"ten_dang_nhap": ten_dang_nhap, "mat_khau": MAT_KHAU}
    )
    assert resp.status_code == 303


# ------------------------- Dự báo -------------------------


def test_du_bao_du_15_xa_va_hotspot_hut_muc_tieu(db):
    du_bao = ban_tin.du_bao_giai_ngan(db)
    assert len(du_bao) == 15
    theo_ma = {d.ma: d for d in du_bao}
    # 2 xã điểm nóng (<30% tháng 7) chắc chắn dự báo hụt mục tiêu 95%
    assert not theo_ma["CACSON"].dat_muc_tieu
    assert not theo_ma["MUONGLAT"].dat_muc_tieu
    # Dự báo nằm trong [0; 100] và xã hụt phải có mức cần tăng tốc > 0
    for d in du_bao:
        assert 0 <= d.ty_le_du_bao <= 100
        if not d.dat_muc_tieu:
            assert d.can_tang_toc_thang > 0


def test_hoi_quy_tuyen_tinh():
    # Chuỗi tăng đều 10/tháng → hệ số góc ~10
    goc, chan = ban_tin._hoi_quy_tuyen_tinh([(1, 10), (2, 20), (3, 30)])
    assert abs(goc - 10) < 1e-9
    assert abs(chan) < 1e-9


def test_lap_ban_tin_co_du_3_viec(db):
    bt = ban_tin.lap_ban_tin(db)
    assert len(bt["viec_can_chi_dao"]) == 3
    assert bt["xa_hut"]
    assert bt["ky"] == 202607


# ------------------------- Trang bản tin + công văn -------------------------


def test_trang_ban_tin(client):
    dang_nhap(client, "lanhdao")
    resp = client.get("/ban-tin")
    assert resp.status_code == 200
    assert "Dự báo giải ngân đến 31/12" in resp.text
    assert "Ba việc cần chỉ đạo hôm nay" in resp.text
    assert "Xã Các Sơn" in resp.text


def test_ban_tin_phan_quyen(client):
    dang_nhap(client, "xa.hacthanh")
    assert client.get("/ban-tin").status_code == 403


def test_du_thao_cong_van_nd30(client):
    dang_nhap(client, "lanhdao")
    resp = client.get("/ban-tin/du-thao-cong-van")
    assert resp.status_code == 200
    doc = Document(io.BytesIO(resp.content))
    van_ban = "\n".join(p.text for p in doc.paragraphs)
    for bang in doc.tables:
        for hang in bang.rows:
            for o in hang.cells:
                van_ban += "\n" + "\n".join(p.text for p in o.paragraphs)
    assert "V/v đôn đốc" in van_ban
    assert "Kính gửi" in van_ban
    assert "./." in van_ban
    assert "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM" in van_ban


# ------------------------- Hộ chiếu số liệu -------------------------


def test_ho_chieu_so_lieu(client, db):
    dang_nhap(client, "lanhdao")
    dv = db.query(DonVi).filter_by(ma="CACSON").one()
    gt = db.query(GiaTriChiTieu).filter_by(don_vi_id=dv.id, nam=2026, thang=7).first()
    resp = client.get(f"/so-lieu/{gt.id}")
    assert resp.status_code == 200
    assert "Hộ chiếu số liệu" in resp.text
    assert "Xã Các Sơn" in resp.text
    assert "CSDL nguồn theo Danh mục QĐ 2053/QĐ-UBND" in resp.text
    assert client.get("/so-lieu/999999").status_code == 404


def test_ho_chieu_can_dang_nhap(client):
    assert client.get("/so-lieu/1").status_code == 303


# ------------------------- Đồng hồ tiết kiệm -------------------------


def test_dong_ho_tiet_kiem(db):
    dh = dong_ho_tiet_kiem(db)
    assert dh["luot_toan_tinh"] > 5000  # 12 báo cáo × 166 xã × 7 tháng vận hành
    assert dh["gio_cong"] == dh["luot_toan_tinh"] * 4


@pytest.mark.parametrize("duong_dan", ["/", "/cong-khai"])
def test_dong_ho_hien_thi(client, duong_dan):
    if duong_dan == "/":
        dang_nhap(client, "lanhdao")
    resp = client.get(duong_dan)
    assert resp.status_code == 200
    assert "lượt báo cáo" in resp.text
