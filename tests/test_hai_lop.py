"""Kiểm thử v0.2: Lớp 1 (kho văn bản, FTS5, chặn mật) + Kênh 2 (máy trích
xuất – người xác nhận) + AI xuyên 2 lớp."""

import io

import pytest
from docx import Document
from sqlalchemy import text

from app.db import SessionLocal
from app.models import ChiTieu, DonVi, GiaTriChiTieu, NguoiDung, VanBan
from app.services import ai_query, extractor, search, van_ban_mau
from scripts.seed import reset_db, seed_all

MAT_KHAU = "Demo@2026"


@pytest.fixture(scope="module", autouse=True)
def reseed_sau_module():
    yield
    reset_db()
    db = SessionLocal()
    try:
        seed_all(db)
    finally:
        db.close()


def dang_nhap(client, ten):
    assert (
        client.post(
            "/dang-nhap", data={"ten_dang_nhap": ten, "mat_khau": MAT_KHAU}
        ).status_code
        == 303
    )


# ------------------------- Lớp 1: seed + FTS5 + mật -------------------------


def test_seed_lop_1_du_van_ban(db):
    assert db.query(VanBan).count() >= 45
    assert db.query(VanBan).filter_by(mat=True).count() == 1
    so_dong_fts = db.execute(text("SELECT COUNT(*) FROM van_ban_fts")).scalar()
    assert so_dong_fts > 100


def test_van_ban_mat_khong_vao_fts_va_tim_kiem(db):
    """Văn bản mật: không có trong chỉ mục, không ra kết quả tìm kiếm."""
    vb_mat = db.query(VanBan).filter_by(mat=True).one()
    trong_fts = db.execute(
        text("SELECT COUNT(*) FROM van_ban_fts WHERE van_ban_id = :id"),
        {"id": vb_mat.id},
    ).scalar()
    assert trong_fts == 0
    assert search.tim_kiem(db, "bi-mat-demo") == []


def test_tim_kiem_fts_ra_ket_qua_kem_nguon(db):
    kq = search.tim_kiem(db, "giải phóng mặt bằng")
    assert kq, "FTS5 phải tìm thấy kế hoạch của Sở Tài chính"
    assert kq[0]["so_ky_hieu"] and kq[0]["ngay_ban_hanh"]


def test_tim_kiem_loc_theo_quyen_chuyen_vien_xa(db):
    """Chuyên viên xã chỉ thấy văn bản đơn vị mình + văn bản sở/tỉnh."""
    nguoi_dung = db.query(NguoiDung).filter_by(ten_dang_nhap="xa.hacthanh").one()
    duoc_phep = {
        dv.id for dv in db.query(DonVi).filter(DonVi.loai.in_(["so_nganh", "tinh"]))
    } | {nguoi_dung.don_vi_id}
    for kq in search.tim_kiem(db, "an sinh xã hội bảo trợ", nguoi_dung, gioi_han=8):
        vb = db.get(VanBan, kq["van_ban_id"])
        assert vb.co_quan_id in duoc_phep


def test_mat_khong_xem_duoc_voi_vai_tro_thuong(client, db):
    vb_mat = db.query(VanBan).filter_by(mat=True).one()
    dang_nhap(client, "lanhdao")
    assert client.get(f"/van-ban/{vb_mat.id}").status_code == 404
    dang_nhap(client, "admin")
    assert client.get(f"/van-ban/{vb_mat.id}").status_code == 200


def test_trang_kho_van_ban_va_tim_kiem(client):
    dang_nhap(client, "lanhdao")
    resp = client.get("/van-ban")
    assert resp.status_code == 200
    assert "Kho văn bản" in resp.text
    resp = client.get("/van-ban?q=giải phóng mặt bằng")
    assert "Kết quả tìm toàn văn" in resp.text
    assert "KH-STC" in resp.text


# ------------------------- Kênh 2: máy trích xuất -------------------------


def test_extractor_bat_dung_so_tu_cau_mau(db):
    """Extractor phải bắt đúng các ca "45.000 triệu đồng" và "27,0%"."""
    dv = db.query(DonVi).filter_by(ma="NGASON").one()
    so = {
        "DTC01": 45000,
        "DTC02": 21500,
        "DTC03": 47.8,
        "DTC04": 9,
        "DTC05": 1,
        "TTHC01": 1250,
        "TTHC02": 1190,
        "TTHC03": 60,
        "TTHC04": 95.2,
        "TTHC05": 780,
        "TTHC06": 62.4,
        "AS01": 120,
        "AS02": 180,
        "AS03": 640,
        "AS04": 372.5,
        "AS05": 27.0,
    }
    vb = VanBan(
        so="99",
        ky_hieu="BC-TEST",
        loai="bao_cao",
        trich_yeu=van_ban_mau.trich_yeu_bao_cao(7, 2026),
        co_quan_id=dv.id,
        toan_van="\n\n".join(van_ban_mau.noi_dung_bao_cao(dv.ten, 7, 2026, so)),
        mat=False,
    )
    db.add(vb)
    db.commit()

    ds = extractor.trich_xuat_offline(db, vb)
    theo_ma = {d.chi_tieu.ma: d for d in ds}
    assert theo_ma["DTC01"].gia_tri_may_doc == 45000  # "45.000 triệu đồng"
    assert theo_ma["DTC02"].gia_tri_may_doc == 21500
    assert theo_ma["AS05"].gia_tri_may_doc == 27.0  # "27,0%"
    assert theo_ma["TTHC01"].gia_tri_may_doc == 1250  # "1.250 hồ sơ"
    # Kỳ suy từ trích yếu; đơn vị theo cơ quan ban hành
    assert theo_ma["DTC02"].nam == 2026 and theo_ma["DTC02"].thang == 7
    assert theo_ma["DTC02"].don_vi_id == dv.id
    # Câu trích dẫn chứa nguyên văn con số
    assert "21.500" in theo_ma["DTC02"].doan_trich
    # Chỉ tiêu dẫn xuất KHÔNG trích (hệ thống tự tính)
    assert "DTC03" not in theo_ma
    # Dọn: hàng chờ test này bị xóa cùng văn bản khi reseed cuối module


def test_extractor_khong_chay_tren_van_ban_mat(db):
    vb_mat = db.query(VanBan).filter_by(mat=True).one()
    assert extractor.trich_xuat_offline(db, vb_mat) == []


def test_gia_tri_ai_khong_nam_trong_doan_trich_bi_loai():
    """Yêu cầu CLAUDE.md: giá trị không xuất hiện nguyên văn → loại bỏ."""
    assert extractor._gia_tri_nam_trong_doan(21500, "đạt 21.500 triệu đồng")
    assert not extractor._gia_tri_nam_trong_doan(99999, "đạt 21.500 triệu đồng")


# ------------------------- Kênh 2: luồng tải lên → xác nhận -------------------------


def _tao_file_docx_bao_cao(db) -> bytes:
    """Dựng file .docx báo cáo tháng 7 của Hạc Thành trong bộ nhớ."""
    dv = db.query(DonVi).filter_by(ma="HACTHANH").one()
    so = {
        "DTC01": 81888,
        "DTC02": 45000,
        "DTC03": 55.0,
        "DTC04": 14,
        "DTC05": 1,
        "TTHC01": 1250,
        "TTHC02": 1200,
        "TTHC03": 50,
        "TTHC04": 96.0,
        "TTHC05": 800,
        "TTHC06": 64.0,
        "AS01": 40,
        "AS02": 55,
        "AS03": 700,
        "AS04": 410.5,
        "AS05": 58.0,
    }
    doc = Document()
    doc.add_paragraph("UBND PHƯỜNG HẠC THÀNH")
    doc.add_paragraph("Số: 58/BC-UBND")
    doc.add_paragraph("BÁO CÁO")
    doc.add_paragraph(van_ban_mau.trich_yeu_bao_cao(7, 2026))
    doc.add_paragraph("Hạc Thành, ngày 24 tháng 7 năm 2026")
    for doan in van_ban_mau.noi_dung_bao_cao(dv.ten, 7, 2026, so):
        doc.add_paragraph(doan)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_luong_tai_len_trich_xuat_xac_nhan(client, db):
    """KỊCH BẢN DEMO TRỌN VẸN: tải báo cáo → máy trích → xác nhận độ tin cậy
    cao → số vào Lớp 2 với nguồn văn bản + liên kết văn bản gốc."""
    dang_nhap(client, "xa.hacthanh")

    # Bước 1: tải file lên — hệ thống bóc siêu dữ liệu, cho rà lại
    noi_dung = _tao_file_docx_bao_cao(db)
    resp = client.post(
        "/van-ban/tai-len",
        files={
            "tep": (
                "bao-cao-thang7.docx",
                noi_dung,
                "application/vnd.openxmlformats-officedocument"
                ".wordprocessingml.document",
            )
        },
    )
    assert resp.status_code == 200
    assert "Rà siêu dữ liệu" in resp.text
    assert "58" in resp.text  # số văn bản bóc được
    ten_tam = resp.text.split('name="ten_tam" value="')[1].split('"')[0]

    # Bước 2: lưu vào Kho → máy trích xuất chạy → chuyển đến hàng chờ
    dv = db.query(DonVi).filter_by(ma="HACTHANH").one()
    resp = client.post(
        "/van-ban/luu",
        data={
            "ten_tam": ten_tam,
            "so": "58",
            "ky_hieu": "BC-UBND",
            "loai": "bao_cao",
            "trich_yeu": van_ban_mau.trich_yeu_bao_cao(7, 2026),
            "co_quan_id": str(dv.id),
            "ngay_ban_hanh": "2026-07-24",
        },
    )
    assert resp.status_code == 303
    assert "/trich-xuat" in resp.headers["location"]

    # Bước 3: hàng chờ hiển thị số máy đọc kèm câu trích + link văn bản gốc
    resp = client.get(resp.headers["location"])
    assert "không phải nhập lại" in resp.text
    assert "DTC02" in resp.text
    assert "Xác nhận tất cả" in resp.text

    # Bước 4: xác nhận tất cả dòng độ tin cậy cao
    resp = client.post("/trich-xuat/xac-nhan-tin-cay-cao")
    assert resp.status_code == 303

    # Kiểm chứng Lớp 2: DTC02 đã vào với nguồn văn bản + liên kết văn bản gốc
    db.expire_all()
    ct = db.query(ChiTieu).filter_by(ma="DTC02").one()
    gt = (
        db.query(GiaTriChiTieu)
        .filter_by(chi_tieu_id=ct.id, don_vi_id=dv.id, nam=2026, thang=7)
        .one()
    )
    assert gt.gia_tri == 45000
    assert gt.nguon == "van_ban"
    assert gt.van_ban_id is not None
    assert gt.nguoi_xac_nhan_id is not None
    # DTC03 được hệ thống tự tính lại từ DTC02/DTC01
    ct3 = db.query(ChiTieu).filter_by(ma="DTC03").one()
    gt3 = (
        db.query(GiaTriChiTieu)
        .filter_by(chi_tieu_id=ct3.id, don_vi_id=dv.id, nam=2026, thang=7)
        .first()
    )
    assert gt3 is not None

    # Hộ chiếu số liệu dẫn về đúng văn bản gốc
    dang_nhap(client, "lanhdao")
    resp = client.get(f"/so-lieu/{gt.id}")
    assert "Kênh 2" in resp.text
    assert f"/van-ban/{gt.van_ban_id}" in resp.text


def test_xa_khac_khong_xac_nhan_ho(client, db):
    """Chuyên viên xã không thấy/không xác nhận được hàng chờ xã khác."""
    dang_nhap(client, "xa.hacthanh")
    resp = client.get("/trich-xuat")
    assert "Xã Ngọc Lặc" not in resp.text  # hàng chờ tồn đọng của Ngọc Lặc


# ------------------------- AI xuyên 2 lớp -------------------------


def test_ai_dinh_tuyen_3_loai():
    assert ai_query.phan_loai_cau_hoi("Những xã nào giải ngân dưới 30%?") == "so_lieu"
    assert (
        ai_query.phan_loai_cau_hoi("Báo cáo nào nói về nguyên nhân chậm?") == "van_ban"
    )
    assert (
        ai_query.phan_loai_cau_hoi(
            "Tỷ lệ đúng hạn là bao nhiêu và có văn bản giải trình chưa?"
        )
        == "lai"
    )


def test_ai_cau_van_ban_tra_doan_kem_nguon(db):
    kq = ai_query.hoi_offline("Báo cáo nào nói về nguyên nhân chậm giải ngân?", db)
    assert kq.loai == "van_ban"
    assert kq.doan_van_ban
    assert kq.doan_van_ban[0]["so_ky_hieu"]  # bắt buộc dẫn nguồn


def test_ai_cau_lai_tra_ca_so_lieu_va_van_ban(db):
    kq = ai_query.hoi_offline(
        "Tỷ lệ đúng hạn TTHC của phường Hạc Thành hiện là bao nhiêu và đã có "
        "báo cáo nào giải trình chưa?",
        db,
    )
    assert kq.loai == "lai"
    assert kq.dong  # có bảng số liệu
    assert kq.doan_van_ban  # có trích đoạn văn bản


def test_ai_khong_lo_van_ban_mat(db):
    kq = ai_query.hoi_offline("Văn bản nào đề cập bi-mat-demo?", db)
    for doan in kq.doan_van_ban:
        assert "bi-mat-demo" not in doan["doan"]


def test_giao_dien_hoi_dap_hien_van_ban_dan_nguon(client):
    dang_nhap(client, "lanhdao")
    resp = client.post(
        "/hoi-dap",
        data={"cau_hoi": "Báo cáo nào nói về nguyên nhân chậm giải ngân?"},
    )
    assert resp.status_code == 200
    assert "Văn bản dẫn nguồn" in resp.text
    assert "lớp văn bản" in resp.text
