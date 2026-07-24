"""Kiểm thử M4: máy soạn báo cáo NĐ30 — file .docx mở được, đúng thể thức."""

import io

import pytest
from docx import Document

from app.models import DonVi
from app.services import report_builder

MAT_KHAU = "Demo@2026"


def dang_nhap(client, ten_dang_nhap):
    resp = client.post(
        "/dang-nhap", data={"ten_dang_nhap": ten_dang_nhap, "mat_khau": MAT_KHAU}
    )
    assert resp.status_code == 303


def van_ban_cua(doc: Document) -> str:
    """Gom toàn bộ chữ trong file .docx (đoạn văn + bảng)."""
    cac_doan = [p.text for p in doc.paragraphs]
    for bang in doc.tables:
        for hang in bang.rows:
            for o in hang.cells:
                cac_doan.extend(p.text for p in o.paragraphs)
    return "\n".join(cac_doan)


@pytest.mark.parametrize("ma_mau", ["BC-DTC", "BC-TTHC"])
def test_sinh_bao_cao_dung_the_thuc(db, ma_mau):
    """python-docx đọc lại được file; có "BÁO CÁO", "./." và quốc hiệu."""
    dv = db.query(DonVi).filter_by(ma="CACSON").one()
    duong_dan = report_builder.tao_bao_cao(db, dv, ma_mau, thang=7)
    assert duong_dan.is_file()

    doc = Document(str(duong_dan))
    van_ban = van_ban_cua(doc)
    assert "BÁO CÁO" in van_ban
    assert "./." in van_ban
    assert "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM" in van_ban
    assert "Độc lập - Tự do - Hạnh phúc" in van_ban
    assert "Nơi nhận:" in van_ban
    assert "TM. ỦY BAN NHÂN DÂN" in van_ban
    assert "Dữ liệu mô phỏng" in van_ban
    assert "Xã Các Sơn".upper() in van_ban.upper()


def test_bao_cao_giai_ngan_co_so_lieu_va_xep_hang(db):
    """Nội dung báo cáo chứa số liệu thật từ CSDL và câu xếp hạng."""
    dv = db.query(DonVi).filter_by(ma="NGASON").one()
    duong_dan = report_builder.tao_bao_cao_giai_ngan(db, dv, thang=7)
    van_ban = van_ban_cua(Document(str(duong_dan)))
    assert "kế hoạch" in van_ban.lower()
    assert "xếp thứ" in van_ban
    assert "Kho dữ liệu dùng chung" in van_ban


def test_kho_du_lieu_thieu_van_sinh_duoc(db):
    """Hạc Thành thiếu DTC02 tháng 7 → báo cáo vẫn sinh, ghi rõ thiếu số liệu."""
    dv = db.query(DonVi).filter_by(ma="HACTHANH").one()
    duong_dan = report_builder.tao_bao_cao_giai_ngan(db, dv, thang=7)
    van_ban = van_ban_cua(Document(str(duong_dan)))
    assert "chưa cập nhật đủ số liệu" in van_ban


def test_api_tao_mot_bao_cao(client):
    dang_nhap(client, "lanhdao")
    resp = client.get("/bao-cao/tao?don_vi=CACSON&mau=BC-DTC&thang=7")
    assert resp.status_code == 200
    assert "wordprocessingml" in resp.headers["content-type"]
    doc = Document(io.BytesIO(resp.content))
    assert "BÁO CÁO" in van_ban_cua(doc)


def test_api_tao_hang_loat_15_xa_kem_thoi_gian(client):
    dang_nhap(client, "lanhdao")
    resp = client.get("/bao-cao/tao-tat-ca?mau=BC-DTC&thang=7")
    assert resp.status_code == 200
    assert "15 báo cáo" in resp.text
    assert "giây" in resp.text


def test_phan_quyen_tao_bao_cao(client):
    dang_nhap(client, "daibieu")
    assert client.get("/bao-cao/tao?don_vi=CACSON&mau=BC-DTC").status_code == 403
    dang_nhap(client, "xa.hacthanh")
    assert client.get("/bao-cao/tao-tat-ca").status_code == 403


def test_tai_file_chan_vuot_thu_muc(client):
    dang_nhap(client, "lanhdao")
    assert client.get("/bao-cao/tai/..%5C.env").status_code in (400, 404)
    assert client.get("/bao-cao/tai/khong-ton-tai.docx").status_code == 404
